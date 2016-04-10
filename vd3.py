from docx import Document

# tìm kiếm và thay thế từ trong word, nhưng còn bị lỗi mất định dạng
#  (do file có sử dụng tô đậm , nghiêng trong đoạn văn)

# mở từ file có sẳn
document = Document("E:\BÁO CÁO NGÀY 01.04.2016 ĐLĐQ - Copy.docx")
# nếu muốn tạo file file mới : document = Document()
a = '÷'
# xác định string unicode
b = u'đến'
for p in document.paragraphs:
    if a in p.text:
        print(p.text)
        text = p.text.replace(a, b)
        style = p.style
        p.text = text
        p.style = style
document.save("E:\BÁO CÁO NGÀY 01.04.2016 ĐLĐQ - Copy.docx")
# nếu save file mới, thì đặt tên hoặc đường dẫn khác.
# Importing modules
